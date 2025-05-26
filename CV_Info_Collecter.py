from tkinter import Tk, Frame, Menu, Label, Button, Toplevel, Entry, PhotoImage # Library
from tkinter import messagebox, filedialog, Checkbutton, IntVar, DISABLED, NORMAL, Radiobutton, colorchooser
import csv # Library
import os # Library
from PIL import ImageTk, Image
import hashlib # Library
import sqlite3 # Library
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class CV_Generator(Frame):
    def __init__(self, master):
        super().__init__(master)
        self.master = master
        self.initUI() # Calls main screen

    def initUI(self):
        # Clear all widgets before adding new UI elements
        for widget in self.master.winfo_children():
            if not isinstance(widget, Menu):  # Keep the menu bar intact
                widget.destroy()  # Destroys the window before it

        self.master.geometry("300x300")
        self.master.configure(bg='white')

        Label(self.master, text="CV Generator", font=("Arial", 20)).pack(fill="x")  # Application name

        cv_path = "cv.csv"
        if not os.path.exists(cv_path):  # If it doesn't exist
            with open(cv_path, "w", encoding="utf-8") as f:  # Creates the file
                f.write("Full Name, Age, County/City, Phone Number, Email, Ethnicity, Education, Skills, Work Experience, Professional Summary, Certifications, Previous Job Title, LinkedIn, Portfolio URL, Extra Info\n")

        # Create the buttons
        self.create_button = Button(self.master, text="Create CV Details", font=("Arial", 12), background="pink", command=lambda: self.collect_cv_info())
        self.create_button.pack(fill="x")

        self.generate_button = Button(self.master, text="Generate Saved CV Details", font=("Arial", 12), background="pink", command=lambda: self.open_collected_cv_info())
        self.generate_button.pack(fill="x")

        # Bind hover effects to buttons for glowing effect
        self.create_button.bind("<Enter>", lambda e: self.create_button.config(bg="yellow"))
        self.create_button.bind("<Leave>", lambda e: self.create_button.config(bg="pink"))

        self.generate_button.bind("<Enter>", lambda e: self.generate_button.config(bg="yellow"))
        self.generate_button.bind("<Leave>", lambda e: self.generate_button.config(bg="pink"))


    def collect_cv_info(self):
        """Collects CV info"""
        """Collects CV info"""
        for widget in self.master.winfo_children():
            if not isinstance(widget, Menu):  # Keep the menu bar intact
                widget.destroy()

        self.master.geometry("470x500")  # Set the desired width x height
        self.master.configure(bg='white')
        
        

        

        def only_numbers(char):
            return char.isdigit()

        vcmd = (self.master.register(only_numbers), '%S')  # %S is the character being inserted

        
        

        Label(self.master, text="Write Your CV Info", font=("Arial", 14)).grid(row=1,column=1)
        
        Label(self.master, text="Your Full Name*", font=("Arial", 9) ,fg="RED").grid(row=2, column=1)
        your_name_entry = Entry(self.master)
        your_name_entry.grid(row=3, column=1)

        Label(self.master, text="Age*", font=("Arial", 9) ,fg="RED").grid(row=4, column=1)
        your_age_entry = Entry(self.master, validate="key", validatecommand=vcmd)
        your_age_entry.grid(row=5, column=1)

        # County/City Entry (no validation needed)
        Label(self.master, text="County/City*", font=("Arial", 9) ,fg="RED").grid(row=6, column=1)
        your_city_entry = Entry(self.master)
        your_city_entry.grid(row=7, column=1)

        # Phone Number Entry
        Label(self.master, text="Phone Number*", font=("Arial", 9) ,fg="RED").grid(row=8, column=1)
        your_number_entry = Entry(self.master, validate="key", validatecommand=vcmd)
        your_number_entry.grid(row=9, column=1)

        Label(self.master, text="Email*", font=("Arial", 9) ,fg="RED").grid(row=10, column=1)
        your_email_entry = Entry(self.master)
        your_email_entry.grid(row=11, column=1)

        Label(self.master, text="Ethnicity*", font=("Arial", 9) ,fg="RED").grid(row=12, column=1)
        your_ethnicity_entry = Entry(self.master)
        your_ethnicity_entry.grid(row=13, column=1)

        Label(self.master, text="Education(Where you got it from/Degree?)*", font=("Arial", 9) ,fg="RED").grid(row=14, column=1)
        your_education_entry = Entry(self.master)
        your_education_entry.grid(row=15, column=1)

        Label(self.master, text="Skill(s)*", font=("Arial", 9) ,fg="RED").grid(row=16, column=1)
        your_skill_entry = Entry(self.master)
        your_skill_entry.grid(row=17, column=1)

        Label(self.master, text="Work Experience*", font=("Arial", 9) ,fg="RED").grid(row=18, column=1)
        your_experience_entry = Entry(self.master)
        your_experience_entry.grid(row=19, column=1)

        Label(self.master, text="Professional Summary", font=("Arial", 9)).grid(row=2, column=2)
        your_summary_entry = Entry(self.master)
        your_summary_entry.grid(row=3, column=2)

        Label(self.master, text="Certifications(If any, split with comma',')", font=("Arial", 9)).grid(row=4, column=2)
        your_certifications_entry = Entry(self.master) 
        your_certifications_entry.grid(row=5, column=2)

        Label(self.master, text="Previous Job Title", font=("Arial", 9)).grid(row=6, column=2)
        your_title_entry = Entry(self.master)
        your_title_entry.grid(row=7, column=2)

        Label(self.master, text="LinkedIn", font=("Arial", 9)).grid(row=8, column=2)
        your_linkedin_entry = Entry(self.master)
        your_linkedin_entry.grid(row=9, column=2)

        Label(self.master, text="Portfolio URL", font=("Arial", 9)).grid(row=10, column=2)
        your_portfolio_entry = Entry(self.master)
        your_portfolio_entry.grid(row=11, column=2)

        Label(self.master, text="Extra Info To Be Aware Of", font=("Arial", 9)).grid(row=12, column=2)
        your_extra_entry = Entry(self.master)
        your_extra_entry.grid(row=13, column=2)

        self.input =  Button(self.master, text="Input Data", background="pink", command=lambda: self.manage_cv_info(
           your_name_entry.get(),
           your_age_entry.get(),
           your_city_entry.get(),
           your_number_entry.get(),
           your_email_entry.get(),
           your_ethnicity_entry.get(),
           your_education_entry.get(),
           your_skill_entry.get(),
           your_experience_entry.get(),
           your_summary_entry.get(),
           your_certifications_entry.get(),
           your_title_entry.get(),
           your_linkedin_entry.get(),
           your_portfolio_entry.get(),
           your_extra_entry.get()
        ))
        self.input.grid(row=400, column=2)


        self.back = Button(self.master, text="Back", background="pink", command=lambda:self.initUI())
        self.back.grid(row=400, column=1)


        self.back.bind("<Enter>", lambda e: self.back.config(bg="yellow"))
        self.back.bind("<Leave>", lambda e: self.back.config(bg="pink"))

        self.input.bind("<Enter>", lambda e: self.input.config(bg="yellow"))
        self.input.bind("<Leave>", lambda e: self.input.config(bg="pink"))


    

    def manage_cv_info(self, name, age, city, number, email, ethnicity, education, skill, experience, summary, certifications, title, linkedin, portfolio, extra):
        """Handle account creation without encryption"""
        print("")
        print("Name is", name)
        print("Age is", age)
        print("City is", city)
        print("Number is", number)
        print("Email is", email)
        print("Ethnicity is", ethnicity)
        print("Education is", education)
        print("Skill(s) is", skill)
        print("Experience is", experience)
        print("Summary is", summary)
        print("Certifications is", certifications)
        print("Previous Job Title is", title)
        print("Linkedin is", linkedin)
        print("Porfolio URL is", portfolio)
        print("Extra Info is", extra)

        missing_fields = []

        if name == "":
            missing_fields.append("Name")
        if age == "":
            missing_fields.append("Age")
        if city == "":
            missing_fields.append("City")
        if number == "":
            missing_fields.append("Phone Number")
        if email == "":
            missing_fields.append("Email")
        if ethnicity == "":
            missing_fields.append("Ethnicity")
        if education == "":
            missing_fields.append("Education")
        if skill == "":
            missing_fields.append("Skill")
        if experience == "":
            missing_fields.append("Work Experience")

        if missing_fields:
            messagebox.showerror("Missing Fields", f"The following fields are missing:\n{', '.join(missing_fields)}")
            return  # Stop and wait for the user to correct input

        with open('cv.csv', 'a', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            
            
            # Write the data row as a list of values
            writer.writerow([name, age, city, number, email, ethnicity, education, skill, experience, summary, certifications, title, linkedin, portfolio, extra])
            messagebox.showinfo("Complete", f"You have done your CV details. Go to Generate Saved CV Details")

        self.initUI()

        
            
                    
    def open_collected_cv_info(self):
        """Collects CV info"""
        for widget in self.master.winfo_children():
            if not isinstance(widget, Menu):  # Keep the menu bar intact
                widget.destroy()

        self.master.geometry("600x500")  # Set the desired width x height
        self.master.configure(bg='light blue')


        Label(self.master, text="Open A Current CV", font=("Arial", 12), bg='light blue').pack()

        try:
            # Open and read the CSV file
            with open('cv.csv', 'r', newline='', encoding='utf-8') as file:
                reader = csv.reader(file)
                header = next(reader)  # Skip the header row

                # Display the header in the GUI
                Label(self.master, text="Full Name(s)", font=("Arial", 10), bg='light blue').pack()

                # Read and create a button for each "Full Name"
                for row in reader:
                    full_name = row[0]  # Get the first column "Full Name"
                    
                    # Create a button for each Full Name
                    button = Button(self.master, text=full_name, font=("Arial", 8), 
                                    command=lambda r=row: self.collect_info(r), background="black", fg="white")
                    button.pack()
                    
                    # Bind mouse hover events to this button to change color
                    button.bind("<Enter>", lambda e, b=button: b.config(bg="blue"))
                    button.bind("<Leave>", lambda e, b=button: b.config(bg="black"))

        except FileNotFoundError:
            messagebox.showerror("Error", "The CV file does not exist.")

        self.back = Button(self.master, text="Back", command=lambda:self.initUI(), background="pink", fg="black")
        self.back.pack()

        self.back.bind("<Enter>", lambda e, b=self.back: b.config(bg="yellow"))
        self.back.bind("<Leave>", lambda e, b=self.back: b.config(bg="pink"))

        

    def collect_info(self, row):
        '''This function get the info for the cv generator'''
        for widget in self.master.winfo_children():
            if not isinstance(widget, Menu):  # Keep the menu bar intact
                widget.destroy()

        self.master.geometry("500x500")  # Set the desired width x height
        self.master.configure(bg='light blue')

        self.color_choice = IntVar()

        Label(self.master, text="Pick CV Colors", font="12", bg='light blue').grid(row=1, column=2)
        self.contrast_var = IntVar()
        self.contrast_check = Checkbutton(
                self.master, 
                text="Tick for contrast colours", 
                variable=self.contrast_var, 
                command=lambda:self.contrasting_check(row),
                bg='light blue'
                

            )
        self.contrast_check.grid()



        self.option1 = Radiobutton(self.master, variable=self.color_choice, value = 1, text="Blue and Orange", bg="light blue", fg="orange", command=print(self.color_choice.get()))
        self.option2 = Radiobutton(self.master, variable=self.color_choice, value = 2, text="Yellow and Purple", bg="light blue", fg="yellow", command=print(self.color_choice.get()))
        self.option3 = Radiobutton(self.master, variable=self.color_choice, value = 3, text="Green and Red", bg="light blue", fg = "red", command=print(self.color_choice.get()))
        self.option4 = Radiobutton(self.master, variable=self.color_choice, value = 4, text="Lime and Pink", bg="light blue", fg = "Green", command=print(self.color_choice.get()))
        self.option1.grid()
        self.option2.grid()
        self.option3.grid()
        self.option4.grid()

        button = Button(self.master, text = "Select color", command=lambda:choose_color)
        button.grid()


        self.option1.config(state=DISABLED)
        self.option2.config(state=DISABLED)
        self.option3.config(state=DISABLED)
        self.option4.config(state=DISABLED)
        

        def choose_color(self):
            # variable to store hexadecimal code of color
            color_code = colorchooser.askcolor(title ="Choose color") 
            print(color_code)

    def contrasting_check(self, row):
        if self.contrast_var.get() == 1:
            self.option1.config(state=NORMAL)
            self.option2.config(state=NORMAL)
            self.option3.config(state=NORMAL)
            self.option4.config(state=NORMAL)
            print("Checked!")
            self.color_choice.set(0)
        else:
            self.option1.config(state=DISABLED)
            self.option2.config(state=DISABLED)
            self.option3.config(state=DISABLED)
            self.option4.config(state=DISABLED)
            print("Unchecked!")
            self.color_choice.set(0)


        self.back = Button(self.master, text="Back", command=lambda:self.initUI(), background="pink", fg="black")
        self.back.grid(row = 400, column=1, pady = 10, padx=10)

        self.colour_data = Button(self.master, text="Input Data", command=lambda:self.generate_cv(row, self.color_choice.get()), background="pink", fg="black")
        self.colour_data.grid(row = 400, column=2, pady = 10, padx=10)

        self.back.bind("<Enter>", lambda e, b=self.back: b.config(bg="yellow"))
        self.back.bind("<Leave>", lambda e, b=self.back: b.config(bg="pink"))

    def generate_cv(self, row, colour_picked):
        """
        This function creates a Word CV with colorful styling like a real resume.
        It uses a left blue sidebar and right white main section layout.
        It saves a new file every time, adding a number if needed.
        """

        print(colour_picked)

        if colour_picked == 0:
            print("Normal")
            left_colour = '#2E75B6'
            right_colour = '#FFC0CB'
        
        if colour_picked == 1:
            print("Blue and Orange")
            left_colour = '#2E75B6'
            right_colour = '#FF6347'   
            
        if colour_picked == 2:
            print("Purple and Yellow")
            left_colour = '#800080'
            right_colour = '#ffff00' 

        if colour_picked == 3:
            print("Green and Red")
            left_colour = '#008000'
            right_colour = '#ff0000' 

        if colour_picked == 4:
            print("Lime and Pink")
            left_colour = '#ffc0cb'
            right_colour = '#00ff00'
            

        # 📝 STEP 1: Create the Word document
        doc = Document()

        # 🧱 STEP 2: Make a table with 1 row and 2 columns for the layout
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Pt(180)  # Sidebar (left)
        table.columns[1].width = Pt(380)  # Main content (right)

        # 🧩 STEP 3: Get the left and right cells from the table
        left = table.cell(0, 0)
        right = table.cell(0, 1)

        # 🧊 STEP 4: Make the left side BLUE using background color
        def set_cell_bg(cell, color_hex):
            cell._element.get_or_add_tcPr()
            tcPr = cell._element.tcPr
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), color_hex)
            tcPr.append(shd)

        set_cell_bg(left, left_colour)  # Blue color like the sample image
        set_cell_bg(right, right_colour)


        # 🏷️ STEP 5: Define helper functions to add text to the sidebar
        def add_sidebar_heading(text):
            para = left.add_paragraph()
            run = para.add_run(text)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # White
            run.font.size = Pt(12)

        def add_sidebar_item(text):
            para = left.add_paragraph()
            run = para.add_run(text)
            run.font.color.rgb = RGBColor(230, 230, 230)  # Light gray
            run.font.size = Pt(10)

        # 📋 STEP 6: Extract values from the `row` list
        labels = [
            "Full Name", "Age", "County/City", "Phone Number", "Email", "Ethnicity", 
            "Education", "Skills", "Work Experience", "Professional Summary", 
            "Certifications", "Previous Job Title", "LinkedIn", "Portfolio URL", "Extra Info"
        ]
        data = dict(zip(labels, row))  # Match labels with values

        # 📦 STEP 7: Fill the sidebar with contact and skills
        add_sidebar_heading("CONTACT")
        add_sidebar_item(f"Phone Number: {data['Phone Number']}")
        add_sidebar_item(f"Email: {data['Email']}")
        add_sidebar_item(f"County/City: {data['County/City']}")

        add_sidebar_heading("\nEDUCATION")
        add_sidebar_item(data["Education"])

        add_sidebar_heading("\nSKILLS")
        add_sidebar_item(data["Skills"])

        add_sidebar_heading("\nLINKS")
        add_sidebar_item("LinkedIn: " + data["LinkedIn"])
        add_sidebar_item("Portfolio: " + data["Portfolio URL"])

        # ➡️ STEP 8: Define helper functions to add styled text to right side
        def add_main_title(text):
            para = right.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = para.add_run(text)
            run.bold = True
            run.underline = True
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(0, 0, 255)

        def add_main_heading(text):
            para = right.add_paragraph()
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(46, 117, 182)  # Nice blue

        def add_main_text(text):
            para = right.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(10)

        # 🧑‍🎓 STEP 9: Add name and title to the top
        add_main_title(data["Full Name"])
        add_main_heading("PREVIOUS JOB TITLE")
        add_main_text(data["Previous Job Title"])

        # 🧠 STEP 10: Add summary
        add_main_heading("PROFILE")
        add_main_text(data["Professional Summary"])

        # 💼 STEP 11: Work experience
        add_main_heading("WORK EXPERIENCE")
        add_main_text(data["Work Experience"])

        # 🏅 STEP 12: Certifications
        add_main_heading("CERTIFICATIONS")
        add_main_text(data["Certifications"])

        # 🧩 STEP 13: Extra info (optional)
        if data["Extra Info"]:
            add_main_heading("EXTRA INFO")
            add_main_text(data["Extra Info"])

        # 💾 STEP 14: Save file with unique name
        base_name = f"{data['Full Name'].replace(' ', '_')}_CV"
        filename = f"{base_name}.docx"
        counter = 1
        while os.path.exists(filename):
            filename = f"{base_name}_{counter}.docx"
            counter += 1

        doc.save(filename)
        print(f"✅ CV saved as: {filename}")

            
        def on_enter(self, e):
            """This function runs when the mouse enters the button."""
            self.cv_button.config(bg="lightblue")  # Change to glowing color (light blue)

        def on_leave(self, e):
            """This function runs when the mouse leaves the button."""
            self.cv_button.config(bg="blue")  # Reset back to the original color

        def open_cv(self):
            """Function to open the CV file when the button is clicked."""
            print("Opening CV...")
                


def main():
    root = Tk()  # Use Tk class here
    root.geometry("300x300")
    app = CV_Generator(root)
    root.mainloop()

if __name__ == '__main__':
    main()